"""Génère le rapport Phase 2 du projet TWM en DOCX."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ── Page de garde ──
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Projet TWM — SFMC Bénin")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Rapport de développement — Phases 1 & 2")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x4B, 0x5B, 0x6B)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Repository : ").bold = True
info.add_run("github.com/Adinette/tp_twm\n")
info.add_run("Auteur : ").bold = True
info.add_run("Freddy ANATO (ak4F)\n")
info.add_run("Date : ").bold = True
info.add_run(f"{date.today().strftime('%d/%m/%Y')}\n")
info.add_run("Technologies : ").bold = True
info.add_run("Next.js 16 · Prisma 7 · NextAuth · PostgreSQL 18 · Tailwind CSS 4")

doc.add_page_break()

# ── Table des matières (manuelle) ──
doc.add_heading("Table des matières", level=1)
toc_items = [
    "1. Contexte du projet",
    "2. Stack technique",
    "3. Phase 1 — Corrections et stabilisation",
    "4. Phase 2 — Inscription et nettoyage",
    "5. Description de la Pull Request",
    "6. Fichiers modifiés (diff)",
    "7. Architecture des routes",
    "8. Schéma de la base de données",
    "9. Tests effectués",
    "10. Prochaines étapes",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ── 1. Contexte ──
doc.add_heading("1. Contexte du projet", level=1)
doc.add_paragraph(
    "Le projet tp_twm est une application web développée dans le cadre du cours "
    "Technologie Web et Mobile à la SFMC Bénin. Le dépôt initial, créé par Adinette, "
    "contenait une base Next.js avec Prisma et NextAuth configurés pour une application "
    "Laravel (tables inutilisées, schéma inadapté). "
    "L'objectif est de transformer ce socle en une application Next.js fonctionnelle "
    "avec un système d'authentification complet."
)

# ── 2. Stack technique ──
doc.add_heading("2. Stack technique", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Technologie"
hdr[1].text = "Version"
hdr[2].text = "Rôle"
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

stack = [
    ("Next.js", "16.2.2", "Framework React full-stack (App Router, Turbopack)"),
    ("React", "19.2.4", "Bibliothèque UI"),
    ("Prisma", "7.7.0", "ORM avec driver adapter pattern"),
    ("NextAuth", "4.24.13", "Authentification (Credentials + Google OAuth)"),
    ("PostgreSQL", "18", "Base de données relationnelle"),
    ("bcrypt", "6.0.0", "Hashage des mots de passe"),
    ("Tailwind CSS", "4", "Framework CSS utilitaire"),
    ("TypeScript", "5", "Typage statique"),
    ("@prisma/adapter-pg", "7.7.0", "Adapter PostgreSQL pour Prisma 7"),
]
for tech, ver, role in stack:
    row = table.add_row().cells
    row[0].text = tech
    row[1].text = ver
    row[2].text = role

# ── 3. Phase 1 ──
doc.add_heading("3. Phase 1 — Corrections et stabilisation", level=1)
doc.add_heading("Branche : fix/prisma-auth-setup", level=2)
doc.add_paragraph("9 fichiers modifiés · 220 insertions · 27 suppressions")

doc.add_heading("3.1 Problèmes identifiés et résolus", level=2)
problems = [
    ("PrismaClient Prisma 7", 
     "L'initialisation classique (datasourceUrl, datasources) ne fonctionne plus en Prisma 7. "
     "Solution : utilisation du driver adapter pattern avec @prisma/adapter-pg et Pool de pg."),
    ("Route NextAuth", 
     "Le dossier [nextauth] a été renommé en [...nextauth] (catch-all route requise par NextAuth)."),
    ("Page de connexion", 
     "Remplacement du bouton générique signIn() par un vrai formulaire (email, mot de passe, "
     "gestion d'erreur, redirection vers /dashboard)."),
    ("SessionProvider manquant", 
     "Création de app/providers.tsx et encapsulation du layout avec <Providers>."),
    ("Middleware de protection", 
     "Correction de la redirection de /login vers /front/auth/login."),
    ("Dashboard", 
     "Création d'une page /dashboard avec affichage de session et bouton de déconnexion."),
]
for title_p, desc in problems:
    p = doc.add_paragraph()
    p.add_run(f"• {title_p} : ").bold = True
    p.add_run(desc)

# ── 4. Phase 2 ──
doc.add_heading("4. Phase 2 — Inscription et nettoyage", level=1)
doc.add_heading("Branche : feature/phase2-register-cleanup", level=2)
doc.add_paragraph("5 fichiers modifiés · 194 insertions · 99 suppressions")

doc.add_heading("4.1 Nettoyage du schéma Prisma", level=2)
doc.add_paragraph(
    "Suppression de 7 tables héritées de Laravel qui n'étaient pas utilisées :"
)
removed = [
    "cache", "cache_locks", "failed_jobs", "job_batches",
    "jobs", "migrations", "sessions", "personal_access_tokens"
]
for t in removed:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph(
    "Simplification du modèle users : id changé de String (dbgenerated) en Int (autoincrement), "
    "ajout de @updatedAt sur updated_at, valeur par défaut now() sur created_at."
)

doc.add_heading("4.2 Système d'inscription", level=2)
doc.add_paragraph("API POST /api/auth/register :")
api_features = [
    "Validation des champs requis (name, email, password)",
    "Vérification de la longueur du mot de passe (≥ 6 caractères)",
    "Contrôle d'unicité de l'email",
    "Hashage du mot de passe avec bcrypt (10 rounds)",
    "Création de l'utilisateur via Prisma",
    "Retour 201 avec message de succès",
    "Gestion des erreurs : 400 (validation), 409 (email existant), 500 (serveur)",
]
for f in api_features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("4.3 Page d'inscription", level=2)
doc.add_paragraph(
    "Création de /front/auth/register avec formulaire complet : nom, email, "
    "mot de passe, confirmation du mot de passe. Validation côté client de la "
    "correspondance des mots de passe. Redirection vers la page de connexion "
    "avec paramètre ?registered=true après succès."
)

doc.add_heading("4.4 Améliorations page de connexion", level=2)
login_improvements = [
    "Bannière verte de succès affichée quand l'utilisateur vient de s'inscrire",
    "Lien « Pas de compte ? S'inscrire » vers la page d'inscription",
    "Navigation bidirectionnelle Login ↔ Register",
]
for imp in login_improvements:
    doc.add_paragraph(imp, style="List Bullet")

doc.add_heading("4.5 Correction import hello/route.ts", level=2)
doc.add_paragraph(
    "Correction du chemin d'import de authOptions : "
    "\"../auth/[nextauth]/route\" → \"../auth/[...nextauth]/route\"."
)

# ── 5. Description PR ──
doc.add_heading("5. Description de la Pull Request", level=1)
doc.add_paragraph()
pr_title = doc.add_paragraph()
pr_title.add_run("Titre : ").bold = True
pr_title.add_run("feat: Phase 2 - Inscription, nettoyage schéma, liens login/register")

doc.add_paragraph()
doc.add_heading("Description complète pour GitHub :", level=2)

pr_desc = """## 🎯 Objectif

Cette PR implémente la Phase 2 du projet : système d'inscription utilisateur et nettoyage du schéma de base de données.

## 📋 Changements

### Nettoyage du schéma Prisma
- Suppression de **7 tables Laravel inutilisées** : `cache`, `cache_locks`, `failed_jobs`, `job_batches`, `jobs`, `migrations`, `sessions`, `personal_access_tokens`
- Modèle `users` simplifié : `id` changé de `String` (dbgenerated) en `Int` (autoincrement), ajout `@updatedAt`

### Système d'inscription
- **API** `POST /api/auth/register` : validation des champs, vérification unicité email, hashage bcrypt, création utilisateur
- **Page** `/front/auth/register` : formulaire complet (nom, email, mot de passe, confirmation)

### Améliorations login
- Bannière de succès après inscription (`?registered=true`)
- Lien « Pas de compte ? S'inscrire » vers `/front/auth/register`

### Corrections
- Fix import `hello/route.ts` : `[...nextauth]` au lieu de `[nextauth]`

## 📊 Statistiques
- **5 fichiers modifiés** · 194 insertions · 99 suppressions

## ✅ Tests
- [x] Inscription : status 201 — utilisateur créé avec email unique
- [x] Connexion : status 302 — redirection après auth réussie
- [x] Providers NextAuth : google + credentials OK
- [x] Bannière succès affichée après inscription
- [x] Navigation Login ↔ Register fonctionnelle

## 🔗 Dépendances
- Basée sur la branche `fix/prisma-auth-setup` (Phase 1)
- Nécessite `npx prisma db push --force-reset` puis `npx prisma generate` après merge"""

doc.add_paragraph(pr_desc)

# ── 6. Fichiers modifiés ──
doc.add_heading("6. Fichiers modifiés (diff complet)", level=1)
doc.add_paragraph("13 fichiers au total entre main et feature/phase2-register-cleanup (Phases 1 + 2) :")

files_table = doc.add_table(rows=1, cols=3)
files_table.style = "Light Grid Accent 1"
files_table.alignment = WD_TABLE_ALIGNMENT.CENTER
fh = files_table.rows[0].cells
fh[0].text = "Fichier"
fh[1].text = "Statut"
fh[2].text = "Description"
for c in fh:
    for p in c.paragraphs:
        p.runs[0].bold = True

files_data = [
    ("app/api/auth/[...nextauth]/route.ts", "Renommé", "Catch-all route NextAuth"),
    ("app/api/auth/register/route.ts", "Nouveau", "API d'inscription POST"),
    ("app/api/hello/route.ts", "Modifié", "Fix import [...nextauth]"),
    ("app/dashboard/page.tsx", "Nouveau", "Page dashboard protégée"),
    ("app/front/auth/login/page.tsx", "Modifié", "Formulaire + bannière + lien register"),
    ("app/front/auth/register/page.tsx", "Nouveau", "Page d'inscription"),
    ("app/layout.tsx", "Modifié", "Ajout Providers, lang=fr"),
    ("app/lib/prisma.ts", "Modifié", "Driver adapter pattern Prisma 7"),
    ("app/providers.tsx", "Nouveau", "SessionProvider wrapper"),
    ("middleware.ts", "Modifié", "Redirect → /front/auth/login"),
    ("package.json", "Modifié", "Ajout @prisma/adapter-pg"),
    ("package-lock.json", "Modifié", "Lock dependencies"),
    ("prisma/schema.prisma", "Modifié", "Schéma nettoyé (2 modèles)"),
]
for fname, status, desc in files_data:
    row = files_table.add_row().cells
    row[0].text = fname
    row[1].text = status
    row[2].text = desc

# ── 7. Architecture ──
doc.add_heading("7. Architecture des routes", level=1)

routes_table = doc.add_table(rows=1, cols=3)
routes_table.style = "Light Grid Accent 1"
routes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rh = routes_table.rows[0].cells
rh[0].text = "Route"
rh[1].text = "Méthode"
rh[2].text = "Description"
for c in rh:
    for p in c.paragraphs:
        p.runs[0].bold = True

routes = [
    ("/front/auth/login", "GET", "Page de connexion (email/password + Google)"),
    ("/front/auth/register", "GET", "Page d'inscription"),
    ("/dashboard", "GET", "Dashboard protégé (nécessite session)"),
    ("/api/auth/register", "POST", "API d'inscription"),
    ("/api/auth/[...nextauth]", "GET/POST", "NextAuth endpoints (signin, callback, etc.)"),
    ("/api/hello", "GET", "Endpoint protégé de test"),
]
for route, method, desc in routes:
    row = routes_table.add_row().cells
    row[0].text = route
    row[1].text = method
    row[2].text = desc

# ── 8. Schéma BDD ──
doc.add_heading("8. Schéma de la base de données", level=1)
doc.add_paragraph("Base : apiProjet · PostgreSQL 18 · localhost:5432")

doc.add_heading("Table users", level=2)
users_table = doc.add_table(rows=1, cols=4)
users_table.style = "Light Grid Accent 1"
uh = users_table.rows[0].cells
uh[0].text = "Colonne"
uh[1].text = "Type"
uh[2].text = "Contrainte"
uh[3].text = "Description"
for c in uh:
    for p in c.paragraphs:
        p.runs[0].bold = True

users_cols = [
    ("id", "Int", "PK, autoincrement", "Identifiant unique"),
    ("name", "Varchar(255)", "NOT NULL", "Nom complet"),
    ("email", "Varchar(255)", "UNIQUE, NOT NULL", "Adresse email"),
    ("password", "Varchar(255)", "NOT NULL", "Hash bcrypt"),
    ("email_verified_at", "DateTime", "NULLABLE", "Date vérification email"),
    ("provider", "Varchar(255)", "NULLABLE", "Provider OAuth (google)"),
    ("provider_id", "Varchar(255)", "NULLABLE", "ID du provider OAuth"),
    ("created_at", "DateTime", "DEFAULT now()", "Date de création"),
    ("updated_at", "DateTime", "@updatedAt", "Dernière modification"),
]
for col, typ, cstr, desc in users_cols:
    row = users_table.add_row().cells
    row[0].text = col
    row[1].text = typ
    row[2].text = cstr
    row[3].text = desc

doc.add_heading("Table password_reset_tokens", level=2)
prt_table = doc.add_table(rows=1, cols=4)
prt_table.style = "Light Grid Accent 1"
ph = prt_table.rows[0].cells
ph[0].text = "Colonne"
ph[1].text = "Type"
ph[2].text = "Contrainte"
ph[3].text = "Description"
for c in ph:
    for p in c.paragraphs:
        p.runs[0].bold = True

prt_cols = [
    ("email", "Varchar(255)", "PK", "Email de l'utilisateur"),
    ("token", "Varchar(255)", "NOT NULL", "Token de réinitialisation"),
    ("created_at", "DateTime", "DEFAULT now()", "Date de création"),
]
for col, typ, cstr, desc in prt_cols:
    row = prt_table.add_row().cells
    row[0].text = col
    row[1].text = typ
    row[2].text = cstr
    row[3].text = desc

# ── 9. Tests ──
doc.add_heading("9. Tests effectués", level=1)

tests_table = doc.add_table(rows=1, cols=4)
tests_table.style = "Light Grid Accent 1"
th = tests_table.rows[0].cells
th[0].text = "Test"
th[1].text = "Commande / Action"
th[2].text = "Résultat attendu"
th[3].text = "Statut"
for c in th:
    for p in c.paragraphs:
        p.runs[0].bold = True

tests = [
    ("Inscription API", "POST /api/auth/register (freddy@test.com)", "Status 201, user créé", "✅ OK"),
    ("Email en double", "POST /api/auth/register (même email)", "Status 409, email déjà utilisé", "✅ OK"),
    ("Connexion credentials", "POST /api/auth/callback/credentials", "Status 302, redirect", "✅ OK"),
    ("Providers NextAuth", "GET /api/auth/providers", "google + credentials", "✅ OK"),
    ("Prisma db push", "npx prisma db push --force-reset", "Schema synchronisé", "✅ OK"),
    ("Prisma generate", "npx prisma generate", "Client généré (v7.7.0)", "✅ OK"),
    ("Serveur dev", "npm run dev", "Ready on localhost:3000", "✅ OK"),
]
for test, cmd, expected, status in tests:
    row = tests_table.add_row().cells
    row[0].text = test
    row[1].text = cmd
    row[2].text = expected
    row[3].text = status

# ── 10. Prochaines étapes ──
doc.add_heading("10. Prochaines étapes", level=1)
next_steps = [
    ("Phase 3 — Composants réutilisables", 
     "Créer un dossier app/components/ avec Navbar, Footer, Card, etc."),
    ("Phase 4 — Dashboard enrichi", 
     "Ajouter des statistiques, graphiques, et gestion de profil utilisateur."),
    ("Phase 5 — Google OAuth", 
     "Configurer GOOGLE_CLIENT_ID et GOOGLE_CLIENT_SECRET pour l'authentification OAuth."),
    ("Phase 6 — Reset mot de passe", 
     "Implémenter le flux complet de réinitialisation de mot de passe via email."),
    ("Phase 7 — Déploiement", 
     "Configurer Vercel ou serveur de production, variables d'environnement, domaine."),
]
for step_title, step_desc in next_steps:
    p = doc.add_paragraph()
    p.add_run(f"• {step_title} : ").bold = True
    p.add_run(step_desc)

# ── Sauvegarde ──
output_path = r"C:\Users\HP\Desktop\cours\Tecnologie Web et Mobile\tp_twm\Rapport_Projet_TWM_Phase2.docx"
doc.save(output_path)
print(f"Rapport généré : {output_path}")
