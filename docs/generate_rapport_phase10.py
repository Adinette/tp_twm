from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Rapport_Projet_TWM_Phase10.docx'

PHASE10_BRANCH = 'feature/phase10-idle-session-security'
PHASE10_BASE = 'feature/phase9-collaboration-release-hub'

PHASE10_DELIVERABLES = [
    ['app/lib/auth-session-config.ts', 'Nouveau', 'Centralise la configuration de duree de session, delai d inactivite et fenetre d avertissement.'],
    ['app/api/auth/[...nextauth]/route.ts', 'Modifie', 'Consomme la configuration partagee de session plutot qu une valeur locale en dur.'],
    ['app/providers.tsx', 'Modifie', 'Ajoute la detection d inactivite, le compte a rebours et la modal d avertissement.'],
    ['app/layout.tsx', 'Modifie', 'Injecte la configuration de session/inactivite dans le provider global.'],
    ['app/front/auth/login/page.tsx', 'Modifie', 'Distingue une expiration standard d une deconnexion apres inactivite.'],
    ['docs/generate_rapport_phase10.py', 'Nouveau', 'Generateur du rapport habituel pour la Phase 10.'],
]

VALIDATION_ROWS = [
    ['Lint phase 10', 'npx eslint app/lib/auth-session-config.ts app/api/auth/[...nextauth]/route.ts app/providers.tsx app/layout.tsx app/front/auth/login/page.tsx', 'OK'],
    ['Diagnostics VS Code', 'get_errors sur le slice phase 10', 'OK'],
    ['DOCX', 'python docs/generate_rapport_phase10.py', 'OK - Rapport_Projet_TWM_Phase10.docx genere'],
]

PROBLEM_ROWS = [
    ['Expiration absolue uniquement', 'La phase 9 coupait la session sur la date d expiration sans tenir compte de l activite utilisateur', 'Introduction d un idle timeout cote client avec verification continue de l activite'],
    ['Aucune alerte avant coupure', 'L utilisateur etait renvoye directement au login a l expiration', 'Ajout d une modal d avertissement avec compte a rebours et choix explicite'],
    ['Configuration dispersee', 'La duree de session vivait dans la route auth uniquement', 'Creation d un module partage de configuration de session'],
    ['Retour UX trop pauvre', 'Le login ne differenciait pas une expiration standard d une fermeture pour inactivite', 'Ajout de messages de retour adaptes sur la page de connexion'],
]

PHASE10_HIGHLIGHTS = [
    'Passage d une expiration fixe a une securite orientee inactivite utilisateur.',
    'Ajout d une modal d avertissement avant deconnexion avec action pour rester connecte.',
    'Configuration unifiee des delais de session et d inactivite pour eviter les divergences entre serveur et client.',
]

PHASE10_CHECKLIST = [
    'Configurer au besoin NEXTAUTH_SESSION_MAX_AGE_SECONDS pour la duree maximale de session.',
    'Configurer NEXTAUTH_IDLE_TIMEOUT_SECONDS pour definir la duree d inactivite avant coupure.',
    'Configurer NEXTAUTH_IDLE_WARNING_SECONDS pour definir le delai d avertissement avant deconnexion.',
    'Tester le comportement de la modal en laissant une session inactive puis en cliquant sur Rester connecte.',
]

doc = Document()


def setup_styles() -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_title_page() -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rapport de Projet - Phase 10')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run('Securite de session avancee par inactivite et avertissement avant deconnexion')
    sub.font.size = Pt(15)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Projet : ').bold = True
    meta.add_run('tp_twm | SFMC Benin | Groupe 5\n')
    meta.add_run('Depot GitHub : ').bold = True
    meta.add_run('github.com/Adinette/tp_twm\n')
    meta.add_run('Branche active : ').bold = True
    meta.add_run(f'{PHASE10_BRANCH}\n')
    meta.add_run('Base technique : ').bold = True
    meta.add_run(f'{PHASE10_BASE}\n')
    meta.add_run('Date : ').bold = True
    meta.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()


def add_table(headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = value


def add_bullet(text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build_document() -> None:
    doc.add_heading('Table des matieres', level=1)
    toc_items = [
        '1. Objectif de la phase 10',
        '2. Ce qui change par rapport a la phase 9',
        '3. Ce qui a ete livre',
        '4. Fichiers modifies',
        '5. Validation effectuee',
        '6. Problemes rencontres et solutions',
        '7. Parametrage et mode operatoire',
        '8. Suite recommandee',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading('1. Objectif de la phase 10', level=1)
    doc.add_paragraph(
        'La Phase 10 approfondit la securite de session introduite en Phase 9. '
        'L objectif n est plus seulement de fixer une duree maximale de session, mais de detecter une veritable inactivite utilisateur, '
        'd avertir avant la coupure et de permettre une prolongation explicite de la session.'
    )
    for item in PHASE10_HIGHLIGHTS:
        add_bullet(item)

    doc.add_paragraph()
    doc.add_heading('2. Ce qui change par rapport a la phase 9', level=1)
    add_table(
        ['Sujet', 'Phase 9', 'Phase 10'],
        [
            ['Declencheur', 'Expiration temporelle fixe', 'Inactivite reelle + expiration absolue'],
            ['Experience utilisateur', 'Redirection directe vers le login', 'Modal d avertissement avant deconnexion'],
            ['Configuration', 'Duree de session seule', 'Session max + delai d inactivite + delai d avertissement'],
            ['Retour login', 'Message simple d expiration', 'Message distinct pour expiration et inactivite'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('3. Ce qui a ete livre', level=1)
    doc.add_heading('3.1 Configuration partagee', level=2)
    add_bullet('Ajout de app/lib/auth-session-config.ts pour centraliser les delais de session et d inactivite.')
    add_bullet('Lecture des variables NEXTAUTH_SESSION_MAX_AGE_SECONDS, NEXTAUTH_IDLE_TIMEOUT_SECONDS et NEXTAUTH_IDLE_WARNING_SECONDS.')

    doc.add_heading('3.2 Detection d inactivite', level=2)
    add_bullet('Ecoute des interactions utilisateur principales : clavier, souris, scroll, touch, focus et retour sur onglet visible.')
    add_bullet('Calcul continu du temps restant avant fermeture sur la base de la derniere activite detectee.')

    doc.add_heading('3.3 Avertissement et prolongation', level=2)
    add_bullet('Affichage d une modal lorsque la fenetre d avertissement est atteinte.')
    add_bullet('Ajout d un bouton Rester connecte qui relance l activite locale et tente un refresh de session.')
    add_bullet('Ajout d un bouton Se deconnecter pour fermer explicitement la session sans attendre.')

    doc.add_heading('3.4 Retour utilisateur', level=2)
    add_bullet('La page de connexion distingue maintenant une expiration standard d une fermeture apres inactivite.')

    doc.add_paragraph()
    doc.add_heading('4. Fichiers modifies', level=1)
    add_table(['Fichier', 'Statut', 'Apport'], PHASE10_DELIVERABLES)

    doc.add_paragraph()
    doc.add_heading('5. Validation effectuee', level=1)
    add_table(['Type', 'Commande / outil', 'Resultat'], VALIDATION_ROWS)
    doc.add_paragraph('Commande de validation principale :')
    add_code(
        'npx eslint app/lib/auth-session-config.ts app/api/auth/[...nextauth]/route.ts '
        'app/providers.tsx app/layout.tsx app/front/auth/login/page.tsx'
    )

    doc.add_paragraph()
    doc.add_heading('6. Problemes rencontres et solutions', level=1)
    add_table(['Probleme', 'Cause', 'Solution'], PROBLEM_ROWS)

    doc.add_paragraph()
    doc.add_heading('7. Parametrage et mode operatoire', level=1)
    for item in PHASE10_CHECKLIST:
        add_bullet(item)
    doc.add_paragraph('Exemple de parametrage :')
    add_code(
        'NEXTAUTH_SESSION_MAX_AGE_SECONDS=28800\n'
        'NEXTAUTH_IDLE_TIMEOUT_SECONDS=1800\n'
        'NEXTAUTH_IDLE_WARNING_SECONDS=60'
    )

    doc.add_paragraph()
    doc.add_heading('8. Suite recommandee', level=1)
    add_bullet('Ouvrir une PR Phase 10 apres validation fonctionnelle manuelle du scenario d inactivite.')
    add_bullet('Tester le comportement sur mobile et dans les changements d onglets longs.')
    add_bullet('Si necessaire, specialiser les delais sur certaines zones sensibles du produit.')

    doc.add_paragraph()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'La Phase 10 fait passer la securite de session du projet a un niveau plus realiste. '
        'Au lieu de s appuyer seulement sur une date d expiration, le systeme detecte maintenant une inactivite veritable, '
        'avertit l utilisateur avant coupure et permet une prolongation explicite de la session. '
        'Le resultat est plus sur, plus lisible et plus conforme a un usage applicatif reel.'
    )


def main() -> None:
    setup_styles()
    add_title_page()
    build_document()
    doc.save(OUTPUT)
    print(f'Rapport genere : {OUTPUT}')


if __name__ == '__main__':
    main()