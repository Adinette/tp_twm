from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    h.font.name = 'Calibri'

# ── Fonctions utilitaires ──
def add_colored_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" — {text}")
    else:
        p.add_run(text)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # Fond gris clair
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear'
    })
    shading.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

# ══════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Rapport de Phase 3")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Composants UI, Page d'accueil & Dashboard Sidebar")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

project = doc.add_paragraph()
project.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = project.add_run("Projet TWM — tp_twm")
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"Branche : feature/phase3-components-ui\n")
run.font.size = Pt(11)
run = info.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y')}\n")
run.font.size = Pt(11)
run = info.add_run("Auteur : Freddy ANATO (ak4F)")
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════
add_colored_heading("Table des matières", 1)
toc_items = [
    "1. Contexte et objectifs",
    "2. Architecture des composants",
    "3. Composant Navbar",
    "4. Composant Footer",
    "5. Composant DashboardSidebar",
    "6. Page d'accueil",
    "7. Dashboard — Layout et Auth Guard",
    "8. Dashboard — Pages",
    "9. Layout global",
    "10. Récapitulatif des fichiers",
    "11. Stack technique",
    "12. Tests et validation",
    "13. Conclusion",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. CONTEXTE ET OBJECTIFS
# ══════════════════════════════════════════════════════════════
add_colored_heading("1. Contexte et objectifs", 1)
doc.add_paragraph(
    "La Phase 3 du projet TWM fait suite aux deux phases précédentes :"
)
add_bullet("Correction des bugs Prisma 7, NextAuth et configuration PostgreSQL", "Phase 1")
add_bullet("Nettoyage du schéma, système d'inscription et formulaires d'authentification", "Phase 2")

doc.add_paragraph(
    "L'objectif de cette Phase 3 est de transformer l'application d'un prototype fonctionnel "
    "en une interface utilisateur professionnelle et ergonomique. Les axes principaux sont :"
)
add_bullet("Création de composants UI réutilisables (Navbar, Footer, Sidebar)")
add_bullet("Refonte complète de la page d'accueil (remplacement du boilerplate Next.js)")
add_bullet("Restructuration du dashboard avec sidebar de navigation et sous-pages dédiées")
add_bullet("Intégration d'un auth guard centralisé dans le layout du dashboard")
add_bullet("Responsive design avec support mobile (menu hamburger, sidebar cachée)")

# ══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE DES COMPOSANTS
# ══════════════════════════════════════════════════════════════
add_colored_heading("2. Architecture des composants", 1)
doc.add_paragraph(
    "La Phase 3 introduit une architecture de composants structurée dans le dossier app/components/ "
    "et enrichit le dashboard avec un layout dédié et plusieurs sous-pages."
)

add_code_block(
    "app/\n"
    "├── components/\n"
    "│   ├── Navbar.tsx          (NOUVEAU - 122 lignes)\n"
    "│   ├── Footer.tsx          (NOUVEAU - 64 lignes)\n"
    "│   └── DashboardSidebar.tsx (NOUVEAU - 102 lignes)\n"
    "├── dashboard/\n"
    "│   ├── layout.tsx          (NOUVEAU - 35 lignes)\n"
    "│   ├── page.tsx            (MODIFIÉ - refactorisé)\n"
    "│   ├── profile/page.tsx    (NOUVEAU - 72 lignes)\n"
    "│   ├── services/page.tsx   (NOUVEAU - 72 lignes)\n"
    "│   └── settings/page.tsx   (NOUVEAU - 62 lignes)\n"
    "├── layout.tsx              (MODIFIÉ - +Navbar/Footer)\n"
    "└── page.tsx                (MODIFIÉ - page d'accueil)"
)

# ══════════════════════════════════════════════════════════════
# 3. COMPOSANT NAVBAR
# ══════════════════════════════════════════════════════════════
add_colored_heading("3. Composant Navbar", 1)
doc.add_paragraph(
    "Le composant Navbar (app/components/Navbar.tsx, 122 lignes) est une barre de navigation "
    "responsive intégrée globalement dans le layout racine de l'application."
)

add_colored_heading("Fonctionnalités", 2)
add_bullet("Logo « TWM » avec lien vers la page d'accueil")
add_bullet("Liens de navigation : Accueil, Dashboard")
add_bullet("Gestion de session via useSession() de NextAuth")
add_bullet("Mode connecté : avatar avec initiale, nom de l'utilisateur, bouton Déconnexion")
add_bullet("Mode déconnecté : boutons Connexion et Inscription")
add_bullet("Menu hamburger pour mobile avec toggle useState")
add_bullet("Design : fond blanc, ombres, transitions hover, max-w-7xl centré")

add_colored_heading("Hooks et dépendances", 2)
add_table(
    ["Hook / Import", "Source", "Usage"],
    [
        ["useSession()", "next-auth/react", "Récupérer la session utilisateur"],
        ["signOut()", "next-auth/react", "Déconnexion"],
        ["useState", "react", "Toggle menu mobile"],
        ["Link", "next/link", "Navigation SPA"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 4. COMPOSANT FOOTER
# ══════════════════════════════════════════════════════════════
add_colored_heading("4. Composant Footer", 1)
doc.add_paragraph(
    "Le composant Footer (app/components/Footer.tsx, 64 lignes) est un pied de page "
    "avec une structure en grille 3 colonnes."
)

add_colored_heading("Structure", 2)
add_table(
    ["Colonne", "Contenu"],
    [
        ["Marque", "Logo TWM, description du projet"],
        ["Navigation", "Liens : Accueil, Dashboard, Services"],
        ["Compte", "Liens : Connexion, Inscription, Mon profil"],
    ]
)
doc.add_paragraph(
    "Le footer inclut un copyright dynamique avec l'année courante via new Date().getFullYear()."
)

# ══════════════════════════════════════════════════════════════
# 5. COMPOSANT DASHBOARDSIDEBAR
# ══════════════════════════════════════════════════════════════
add_colored_heading("5. Composant DashboardSidebar", 1)
doc.add_paragraph(
    "Le composant DashboardSidebar (app/components/DashboardSidebar.tsx, 102 lignes) "
    "fournit la navigation latérale dans le dashboard."
)

add_colored_heading("Items de navigation", 2)
add_table(
    ["Label", "Route", "Icône"],
    [
        ["Vue d'ensemble", "/dashboard", "SVG carrés (grille)"],
        ["Mon profil", "/dashboard/profile", "SVG utilisateur"],
        ["Microservices", "/dashboard/services", "SVG serveur"],
        ["Paramètres", "/dashboard/settings", "SVG engrenage"],
    ]
)

add_colored_heading("Fonctionnalités", 2)
add_bullet("État actif dynamique via usePathname() — mise en surbrillance bleue du lien courant")
add_bullet("Affichage du nom et email de l'utilisateur connecté")
add_bullet("Bouton de déconnexion avec signOut({ callbackUrl: '/' })")
add_bullet("Design : fond gris foncé (gray-900), texte blanc, icônes SVG inline")
add_bullet("Responsive : hidden lg:block — masqué sur mobile, visible sur desktop")

# ══════════════════════════════════════════════════════════════
# 6. PAGE D'ACCUEIL
# ══════════════════════════════════════════════════════════════
add_colored_heading("6. Page d'accueil", 1)
doc.add_paragraph(
    "La page d'accueil (app/page.tsx, ~154 lignes) a été entièrement refondue pour "
    "remplacer le boilerplate par défaut de Next.js par une landing page professionnelle."
)

add_colored_heading("Sections", 2)
add_table(
    ["Section", "Description"],
    [
        ["Hero", "Titre « Projet TWM » + sous-titre + 2 CTAs (Commencer → /front/auth/register, Se connecter → /front/auth/login)"],
        ["Features", "3 cartes : Authentification sécurisée, Base de données PostgreSQL, Framework Next.js 16"],
        ["Stack technique", "8 badges technologiques : Next.js, React, TypeScript, Tailwind, Prisma, PostgreSQL, NextAuth, Node.js"],
        ["CTA final", "Section bleue avec incitation à créer un compte + bouton vers /front/auth/register"],
    ]
)

doc.add_paragraph(
    "Le design utilise un dégradé hero (blue-600 → blue-800), des cartes avec ombres et hover effects, "
    "et des badges technologiques en fond bleu-50 avec texte bleu-700."
)

# ══════════════════════════════════════════════════════════════
# 7. DASHBOARD — LAYOUT ET AUTH GUARD
# ══════════════════════════════════════════════════════════════
add_colored_heading("7. Dashboard — Layout et Auth Guard", 1)
doc.add_paragraph(
    "Le fichier app/dashboard/layout.tsx (35 lignes) sert de layout partagé pour toutes les pages "
    "du dashboard. Il centralise la logique d'authentification."
)

add_colored_heading("Fonctionnement", 2)
add_bullet("'use client' — composant côté client pour accéder à useSession()")
add_bullet("useSession() vérifie l'état de la session")
add_bullet("status === 'loading' : affiche un spinner animé plein écran")
add_bullet("status === 'unauthenticated' : redirection automatique vers /front/auth/login via useRouter()")
add_bullet("status === 'authenticated' : affiche DashboardSidebar + zone de contenu ({children})")

add_colored_heading("Avantages", 2)
add_bullet("Supprime la duplication du code d'auth dans chaque page du dashboard")
add_bullet("Garantit qu'aucune page sous /dashboard n'est accessible sans session")
add_bullet("La sidebar est affichée une seule fois pour tout le dashboard")

# ══════════════════════════════════════════════════════════════
# 8. DASHBOARD — PAGES
# ══════════════════════════════════════════════════════════════
add_colored_heading("8. Dashboard — Pages", 1)

add_colored_heading("8.1 Page principale (page.tsx)", 2)
doc.add_paragraph(
    "La page dashboard/page.tsx a été refactorisée : la barre de navigation inline et la logique "
    "d'authentification ont été supprimées (désormais gérées par le layout). Le contenu se compose de :"
)
add_bullet("Message de bienvenue avec le nom de l'utilisateur")
add_bullet("4 cartes de statistiques : Microservices (3), Endpoints (6), PostgreSQL (Connecté), Auth (2 providers)")
add_bullet("4 cartes d'accès rapide avec liens vers : profil, services, API Hello, paramètres")

add_colored_heading("8.2 Page Profil (profile/page.tsx)", 2)
doc.add_paragraph("Page de profil utilisateur (72 lignes) avec :")
add_bullet("Avatar circulaire avec l'initiale du nom")
add_bullet("Affichage du nom et de l'email (lecture seule)")
add_bullet("Section Sécurité : changement de mot de passe (désactivé — « Bientôt disponible »)")
add_bullet("Vérification d'email : badge « En attente »")

add_colored_heading("8.3 Page Microservices (services/page.tsx)", 2)
doc.add_paragraph("Liste des 3 microservices du projet (72 lignes) :")
add_table(
    ["Service", "Endpoint", "Stack"],
    [
        ["Auth Service", "/api/auth/[...nextauth]", "NextAuth, JWT, Prisma"],
        ["API Hello", "/api/hello", "Next.js API Route"],
        ["Register Service", "/api/auth/register", "bcrypt, Prisma, PostgreSQL"],
    ]
)
doc.add_paragraph("Chaque service affiche un badge vert « Actif » et ses informations techniques.")

add_colored_heading("8.4 Page Paramètres (settings/page.tsx)", 2)
doc.add_paragraph("Page de paramètres (62 lignes) avec deux sections :")
add_bullet("Paramètres généraux : Thème (Automatique), Langue (Français)")
add_bullet("Informations techniques : tableau des versions de toutes les technologies du projet")
add_table(
    ["Technologie", "Version"],
    [
        ["Next.js", "16.2.2"],
        ["Prisma", "7.7.0"],
        ["NextAuth", "4.24.13"],
        ["PostgreSQL", "18"],
        ["React", "19.2.4"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 9. LAYOUT GLOBAL
# ══════════════════════════════════════════════════════════════
add_colored_heading("9. Layout global", 1)
doc.add_paragraph(
    "Le fichier app/layout.tsx a été modifié pour intégrer les composants Navbar et Footer "
    "globalement dans toute l'application."
)

add_colored_heading("Modifications", 2)
add_bullet("Import de Navbar depuis ./components/Navbar")
add_bullet("Import de Footer depuis ./components/Footer")
add_bullet("Structure du body : <Providers> → <Navbar /> → <main className='flex-1'>{children}</main> → <Footer /> → </Providers>")
add_bullet("Le body utilise flex flex-col min-h-screen pour que le footer reste en bas")

# ══════════════════════════════════════════════════════════════
# 10. RÉCAPITULATIF DES FICHIERS
# ══════════════════════════════════════════════════════════════
add_colored_heading("10. Récapitulatif des fichiers", 1)
add_table(
    ["Fichier", "Action", "Lignes"],
    [
        ["app/components/Navbar.tsx", "CRÉÉ", "122"],
        ["app/components/Footer.tsx", "CRÉÉ", "64"],
        ["app/components/DashboardSidebar.tsx", "CRÉÉ", "102"],
        ["app/dashboard/layout.tsx", "CRÉÉ", "35"],
        ["app/dashboard/page.tsx", "MODIFIÉ", "~109"],
        ["app/dashboard/profile/page.tsx", "CRÉÉ", "72"],
        ["app/dashboard/services/page.tsx", "CRÉÉ", "72"],
        ["app/dashboard/settings/page.tsx", "CRÉÉ", "62"],
        ["app/layout.tsx", "MODIFIÉ", "+8"],
        ["app/page.tsx", "MODIFIÉ", "~154"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("\nTotal : 10 fichiers · 690 insertions · 110 suppressions")
run.bold = True
run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════
# 11. STACK TECHNIQUE
# ══════════════════════════════════════════════════════════════
add_colored_heading("11. Stack technique", 1)
add_table(
    ["Technologie", "Version", "Rôle"],
    [
        ["Next.js", "16.2.2", "Framework React full-stack (App Router, Turbopack)"],
        ["React", "19.2.4", "Bibliothèque UI"],
        ["TypeScript", "5", "Typage statique"],
        ["Tailwind CSS", "4", "Framework CSS utility-first"],
        ["NextAuth", "4.24.13", "Authentification (Credentials + Google)"],
        ["Prisma", "7.7.0", "ORM avec driver adapter pattern"],
        ["PostgreSQL", "18", "Base de données relationnelle"],
        ["bcrypt", "6.0.0", "Hachage de mots de passe"],
        ["Node.js", "22+", "Runtime JavaScript"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 12. TESTS ET VALIDATION
# ══════════════════════════════════════════════════════════════
add_colored_heading("12. Tests et validation", 1)
doc.add_paragraph("Les vérifications suivantes ont été effectuées :")

add_table(
    ["Test", "Résultat"],
    [
        ["Compilation TypeScript (0 erreurs)", "✅ Validé"],
        ["Navigation Navbar desktop", "✅ Validé"],
        ["Menu hamburger mobile", "✅ Validé"],
        ["Auth guard dashboard (redirect)", "✅ Validé"],
        ["Sidebar état actif", "✅ Validé"],
        ["Page d'accueil (hero, features, stack)", "✅ Validé"],
        ["Dashboard stats et accès rapide", "✅ Validé"],
        ["Page profil (avatar, infos)", "✅ Validé"],
        ["Page microservices (3 services)", "✅ Validé"],
        ["Page paramètres (versions)", "✅ Validé"],
        ["Footer (3 colonnes, copyright)", "✅ Validé"],
        ["Responsive design", "✅ Validé"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 13. CONCLUSION
# ══════════════════════════════════════════════════════════════
add_colored_heading("13. Conclusion", 1)
doc.add_paragraph(
    "La Phase 3 transforme l'application TWM d'un prototype fonctionnel en une interface "
    "utilisateur professionnelle. Les composants Navbar, Footer et DashboardSidebar sont réutilisables "
    "et s'intègrent de manière cohérente dans toute l'application."
)
doc.add_paragraph(
    "Le dashboard est désormais structuré avec un layout dédié (auth guard centralisé), "
    "une sidebar de navigation et 4 sous-pages spécialisées. La page d'accueil présente "
    "le projet de manière attractive avec ses fonctionnalités et sa stack technique."
)
doc.add_paragraph(
    "Cette phase pose les bases d'une expérience utilisateur complète et prépare le terrain "
    "pour les futures fonctionnalités (CRUD microservices, édition de profil, thème sombre, etc.)."
)

# ── Informations Git ──
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Informations Git")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

add_table(
    ["Champ", "Valeur"],
    [
        ["Branche", "feature/phase3-components-ui"],
        ["Base", "feature/phase2-register-cleanup"],
        ["Commit", "2d82d94"],
        ["Fichiers", "10 fichiers modifiés"],
        ["Insertions", "+690"],
        ["Suppressions", "-110"],
        ["Repository", "github.com/Adinette/tp_twm"],
    ]
)

# ── Sauvegarde ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Rapport_Projet_TWM_Phase3.docx")
doc.save(output_path)
print(f"Rapport généré : {output_path}")
