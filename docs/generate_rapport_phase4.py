from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# ── Styles globaux ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    h.font.name = 'Calibri'

# ── Helpers ──
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(f" : {text}")
    else:
        p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shd = p._element.get_or_add_pPr().makeelement(qn('w:shd'),
          {qn('w:fill'): 'F0F0F0', qn('w:val'): 'clear'})
    p._element.get_or_add_pPr().append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return table

def add_alert(text, color_hex='FFF3CD', border_hex='F59E0B'):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠  {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    shd = p._element.get_or_add_pPr().makeelement(qn('w:shd'),
          {qn('w:fill'): color_hex, qn('w:val'): 'clear'})
    p._element.get_or_add_pPr().append(shd)
    return p

# ══════════════════════════════════════════════════════════════
#  PAGE DE TITRE
# ══════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rapport de Projet — Phase 4")
r.bold = True; r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Google OAuth 2.0 — Bilan & Relance collaborative")
r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = proj.add_run("Projet TWM — tp_twm  |  SFMC Bénin  |  Groupe 5")
r.bold = True; r.font.size = Pt(13)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f"Dépôt GitHub : github.com/Adinette/tp_twm\n").font.size = Pt(11)
info.add_run(f"Branche active : feature/phase4-google-oauth\n").font.size = Pt(11)
info.add_run(f"Date : {datetime.now().strftime('%d %B %Y')}\n").font.size = Pt(11)
info.add_run("Rédigé par : Membre 1 (Lead) — à diffuser à tous les membres").font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table des matières", level=1)
toc = [
    "1.  Où en sommes-nous ? — Bilan des 4 phases",
    "2.  Stack technique du projet",
    "3.  Mettre à jour le projet sur sa machine",
    "     3.1  CAS HABITUEL — Déjà cloné lors d'une phase précédente",
    "     3.2  CAS RARE — Premier clone du projet",
    "     3.3  Lancer le projet",
    "4.  Ce qui a été fait en Phase 4 — Google OAuth 2.0",
    "     4.1  Problème résolu : PrismaAdapter incompatible",
    "     4.2  Solution implémentée : callback signIn",
    "     4.3  Obtention des credentials Google Cloud",
    "5.  Tests à effectuer après avoir récupéré le projet",
    "6.  Problèmes connus et solutions",
    "7.  Proposition concrète pour la Phase 5",
    "8.  Répartition des tâches — Équipe de 3 membres",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. BILAN PHASES
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Où en sommes-nous ? — Bilan des 4 phases", level=1)
doc.add_paragraph(
    "Ce tableau résume l'ensemble des travaux réalisés depuis le début du projet. "
    "Il permet à chaque membre de se resituer rapidement."
)
add_table(
    ["Phase", "Branche Git", "Ce qui a été fait", "Statut"],
    [
        ["Phase 1\nCorrections", "fix/prisma-auth-setup",
         "Prisma 7 + driver adapter pg\nNextAuth routing [...nextauth]\nPage login + dashboard\nMiddleware de session",
         "✅ Mergé\nsur main"],
        ["Phase 2\nInscription", "feature/phase2-register-cleanup",
         "Nettoyage schéma Prisma (7 tables supprimées)\nAPI POST /api/auth/register\nPage inscription /front/auth/register\nLien login ↔ register",
         "✅ Poussé\nPR ouverte"],
        ["Phase 3\nUI/Components", "feature/phase3-components-ui",
         "Navbar + Footer responsive\nDashboardSidebar dynamique\nPage d'accueil professionnelle\nDashboard + profil + services + paramètres",
         "✅ Poussé\nPR ouverte"],
        ["Phase 4\nGoogle OAuth", "feature/phase4-google-oauth",
         "Suppression PrismaAdapter (incompatible)\nCallback signIn : création auto user Google\nCredentials Google Cloud configurés\nGuide_Google_OAuth.html",
         "✅ Poussé\nPR à merger"],
    ]
)
doc.add_paragraph()
add_alert("Le déploiement en production (Vercel) n'est pas encore effectué. Il sera traité en Phase 5 après stabilisation du code.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. STACK TECHNIQUE
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Stack technique du projet", level=1)
add_table(
    ["Technologie", "Version", "Rôle dans le projet"],
    [
        ["Next.js", "16.2.2", "Framework principal (App Router, Turbopack)"],
        ["React", "19.2.4", "UI côté client"],
        ["TypeScript", "5", "Typage statique"],
        ["Tailwind CSS", "4", "Styles utilitaires"],
        ["NextAuth.js", "4.24.13", "Auth : Credentials + Google OAuth 2.0 + JWT"],
        ["Prisma", "7.7.0", "ORM — accès PostgreSQL via driver adapter"],
        ["PostgreSQL", "18", "Base de données (locale : apiProjet)"],
        ["bcrypt", "6.0.0", "Hashage mots de passe"],
        ["Google OAuth 2.0", "—", "Connexion sociale via Google Cloud Console"],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  3. PROCÉDURE DE RÉCUPÉRATION DU PROJET
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Mettre à jour le projet sur sa machine", level=1)
doc.add_paragraph(
    "Le dépôt a déjà été cloné lors d'une phase précédente. "
    "Il ne faut donc PAS recloner. Il suffit de récupérer les nouveaux commits "
    "et de basculer sur la bonne branche. Cette section couvre les deux cas possibles."
)

# ── CAS 1 : déjà cloné ──
doc.add_heading("3.1  CAS HABITUEL — Vous avez déjà le projet en local", level=2)
doc.add_paragraph(
    "C'est le cas de tous les membres qui ont participé aux phases précédentes. "
    "Ouvrez un terminal dans le dossier tp_twm et exécutez ces commandes :"
)

doc.add_paragraph("Étape 1 — Sauvegarder d'abord tout travail local en cours :")
add_code(
    "git status                   # voir s'il y a des fichiers modifiés\n"
    "git stash                    # mettre de côté les modifications non commitées (si besoin)"
)

doc.add_paragraph("Étape 2 — Récupérer toutes les nouvelles branches depuis GitHub :")
add_code(
    "git fetch origin             # télécharge sans modifier votre code\n"
    "git branch -a                # vérifier que feature/phase4-google-oauth apparaît"
)

doc.add_paragraph("Étape 3 — Basculer sur la branche Phase 4 :")
add_code(
    "git checkout feature/phase4-google-oauth\n"
    "git pull origin feature/phase4-google-oauth"
)

doc.add_paragraph("Étape 4 — Installer les nouvelles dépendances si le package.json a changé :")
add_code("npm install")

doc.add_paragraph("Étape 5 — Mettre à jour le schéma de base de données :")
add_code("npx prisma db push")
doc.add_paragraph(
    "Cette commande est sans danger : elle applique uniquement les changements "
    "de schéma manquants, sans supprimer vos données existantes."
)

doc.add_paragraph("Étape 6 — Compléter le fichier .env avec les nouvelles clés Google :")
doc.add_paragraph(
    "Ouvrez votre fichier .env existant et ajoutez les deux lignes suivantes "
    "si elles ne sont pas encore présentes (demandez les valeurs au Membre 1) :"
)
add_code(
    "GOOGLE_CLIENT_ID=...   ← demandez la valeur au Membre 1\n"
    "GOOGLE_CLIENT_SECRET=...   ← demandez la valeur au Membre 1"
)
add_alert("Le .env n'est jamais sur GitHub (il est dans .gitignore). Vous devez l'obtenir directement auprès du Membre 1.")

# ── CAS 2 : nouveau clone ──
doc.add_heading("3.2  CAS RARE — Vous n'avez jamais cloné le projet", level=2)
doc.add_paragraph("Uniquement si vous n'avez jamais eu le projet en local :")

add_table(
    ["Outil", "Version minimale", "Téléchargement"],
    [
        ["Node.js", "≥ 20 LTS", "https://nodejs.org"],
        ["Git", "≥ 2.40", "https://git-scm.com"],
        ["PostgreSQL", "≥ 14", "https://www.postgresql.org/download/"],
        ["VS Code", "Dernière version", "https://code.visualstudio.com"],
    ]
)
doc.add_paragraph()
add_code(
    "# 1. Cloner\n"
    "git clone https://github.com/Adinette/tp_twm.git\n"
    "cd tp_twm\n\n"
    "# 2. Basculer sur la branche active\n"
    "git checkout feature/phase4-google-oauth\n\n"
    "# 3. Installer les dépendances\n"
    "npm install\n\n"
    "# 4. Créer le fichier .env (obtenir les valeurs auprès du Membre 1)\n"
    "# 5. Créer la base de données\n"
    "psql -U postgres -c \"CREATE DATABASE \\\"apiProjet\\\";\"\n\n"
    "# 6. Synchroniser le schéma\n"
    "npx prisma db push"
)

doc.add_heading("3.3 Lancer le projet", level=2)
add_code("npm run dev")
doc.add_paragraph("Le projet est accessible sur : http://localhost:3000")
doc.add_paragraph("Pages disponibles :")
add_bullet("http://localhost:3000 → Page d'accueil")
add_bullet("http://localhost:3000/front/auth/login → Connexion (Email/MDP ou Google)")
add_bullet("http://localhost:3000/front/auth/register → Inscription")
add_bullet("http://localhost:3000/dashboard → Tableau de bord (connecté uniquement)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. TRAVAUX RÉALISÉS PHASE 4
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. Ce qui a été fait en Phase 4 — Google OAuth 2.0", level=1)

doc.add_heading("4.1 Problème résolu : PrismaAdapter incompatible", level=2)
doc.add_paragraph(
    "Le code initial utilisait PrismaAdapter pour lier NextAuth à Prisma. "
    "Ce composant exige 4 modèles dans le schéma Prisma (User, Account, Session, VerificationToken) "
    "avec des IDs de type String. Notre schéma utilise un id de type Int, "
    "ce qui provoquait un crash immédiat au démarrage. "
    "Solution : suppression de PrismaAdapter, gestion manuelle via callback."
)

doc.add_heading("4.2 Solution implémentée : callback signIn", level=2)
doc.add_paragraph(
    "À chaque connexion Google, NextAuth appelle signIn. On vérifie si l'utilisateur "
    "existe en base. S'il est nouveau, on le crée automatiquement avec provider='google'."
)
add_code(
    "// app/api/auth/[...nextauth]/route.ts\n"
    "async signIn({ user, account }) {\n"
    '  if (account?.provider === "google") {\n'
    "    const existing = await prisma.users.findUnique({ where: { email: user.email } })\n"
    "    if (!existing) {\n"
    "      await prisma.users.create({\n"
    "        data: {\n"
    "          name: user.name ?? 'Utilisateur Google',\n"
    "          email: user.email,\n"
    '          password: "",\n'
    '          provider: "google",\n'
    "          provider_id: account.providerAccountId,\n"
    "        }\n"
    "      })\n"
    "    }\n"
    "  }\n"
    "  return true\n"
    "}"
)

doc.add_heading("4.3 Obtention des credentials Google Cloud", level=2)
doc.add_paragraph("Les credentials ont été créés sur Google Cloud Console :")
steps = [
    ("URL", "https://console.cloud.google.com"),
    ("Projet créé", "tp-twm-sfmc"),
    ("API activée", "Google People API"),
    ("Type identifiant", "ID client OAuth — Application Web"),
    ("Origine JS autorisée", "http://localhost:3000"),
    ("URI de redirection", "http://localhost:3000/api/auth/callback/google"),
    ("GOOGLE_CLIENT_ID", "Configuré dans .env"),
    ("GOOGLE_CLIENT_SECRET", "Configuré dans .env"),
]
for label, val in steps:
    add_bullet(val, label)

# ══════════════════════════════════════════════════════════════
#  5. TESTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Tests à effectuer après avoir récupéré le projet", level=1)
doc.add_paragraph("Chaque membre doit effectuer ces tests après installation :")
add_table(
    ["#", "Action", "Résultat attendu"],
    [
        ["1", "npm run dev → naviguer sur http://localhost:3000", "Page d'accueil s'affiche"],
        ["2", "Aller sur /front/auth/login", "Formulaire + bouton Google visible"],
        ["3", "S'inscrire via /front/auth/register", "Redirection vers login avec bannière succès"],
        ["4", "Se connecter avec email/mot de passe", "Redirection vers /dashboard"],
        ["5", "Cliquer 'Se connecter avec Google'", "Popup Google → sélection compte"],
        ["6", "Connexion Google réussie", "Redirection /dashboard, nom affiché"],
        ["7", "Vérifier en base : SELECT * FROM users", "User avec provider='google' présent"],
        ["8", "Se déconnecter puis se reconnecter", "Pas de doublon en base"],
    ]
)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  6. PROBLÈMES CONNUS
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Problèmes connus et solutions", level=1)
add_table(
    ["Problème", "Cause", "Solution"],
    [
        ["npm run dev → erreur 'Cannot find module'",
         "npm install non exécuté",
         "Lancer npm install dans le dossier tp_twm"],
        ["Erreur Prisma 'Table users does not exist'",
         "npx prisma db push non exécuté",
         "Lancer npx prisma db push"],
        ["Google OAuth ne redirige pas",
         "GOOGLE_CLIENT_ID ou SECRET vide dans .env",
         "Demander les valeurs au Membre 1 et relancer npm run dev"],
        ["Avertissement 'middleware deprecated'",
         "Next.js 16 renomme middleware.ts en proxy.ts",
         "Non bloquant — à corriger en Phase 5"],
        ["Erreur 'Database connection failed'",
         "PostgreSQL non démarré ou DATABASE_URL incorrect",
         "Démarrer le service PostgreSQL et vérifier le mot de passe dans .env"],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  7. PROPOSITION POUR LA SUITE
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Proposition concrète pour la Phase 5", level=1)
doc.add_paragraph(
    "Maintenant que l'authentification est complète (Credentials + Google OAuth), "
    "la priorité absolue est l'architecture microservices avec Kong Gateway. "
    "C'est le cœur du projet SFMC Bénin. Voici le plan concret :"
)

doc.add_heading("Objectif principal : Kong Gateway + 1er microservice", level=2)
doc.add_paragraph(
    "Kong Gateway est un API Gateway open source. Il centralise les requêtes HTTP, "
    "gère l'authentification, le rate limiting et le routage vers les microservices."
)

doc.add_heading("Étapes concrètes de la Phase 5", level=2)
phase5_steps = [
    ("Étape 1", "Créer docker-compose.yml avec Kong + PostgreSQL pour Kong"),
    ("Étape 2", "Démarrer Kong en local : docker compose up -d"),
    ("Étape 3", "Créer un microservice simple : API transactions (Node.js/Express, port 4000)"),
    ("Étape 4", "Enregistrer le service dans Kong Admin API (port 8001)"),
    ("Étape 5", "Configurer une route Kong : GET /transactions → microservice"),
    ("Étape 6", "Ajouter plugin JWT Kong pour sécuriser les routes"),
    ("Étape 7", "Consommer le service depuis Next.js (page /dashboard/services)"),
    ("Étape 8", "Commit + PR + rapport Phase 5"),
]
for label, val in phase5_steps:
    add_bullet(val, label)

doc.add_paragraph()
doc.add_heading("Après Phase 5 : Déploiement en production", level=2)
add_bullet("Vercel : déploiement Next.js (gratuit, connecté au GitHub repo)")
add_bullet("Base de données cloud : Neon.tech ou Railway (PostgreSQL gratuit)")
add_bullet("Variables .env à configurer dans Vercel Dashboard")
add_bullet("URI Google OAuth à mettre à jour pour le domaine de production")
add_alert("Le déploiement Vercel ne peut se faire que APRÈS que le code soit stable en local sur main.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  8. RÉPARTITION 3 MEMBRES
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Répartition des tâches — Équipe de 3 membres", level=1)
doc.add_paragraph(
    "L'équipe compte 3 membres. La répartition ci-dessous tient compte de cette réalité "
    "et distribue équitablement les tâches selon les profils. Chaque membre a une branche "
    "Git dédiée et ouvre une PR avant de merger."
)

doc.add_heading("Répartition par membre", level=2)
add_table(
    ["Membre", "Rôle", "Phase 5 — Tâche", "Phase 6 — Tâche"],
    [
        ["Membre 1\n(Lead / Full-stack)",
         "Architecture + coordination",
         "docker-compose.yml Kong Gateway\nEnregistrement service + routes Kong\nIntégration Next.js ↔ Kong",
         "Déploiement Vercel + Railway\nVariables env production\nMise à jour URI Google OAuth prod"],
        ["Membre 2\n(Backend)",
         "API / services",
         "Microservice transactions (Node.js/Express)\nCRUD : POST/GET/DELETE /transactions\nConnexion à PostgreSQL",
         "Plugin JWT Kong sur les routes\nMiddleware rôles (ADMIN/USER)\nTests API (Jest)"],
        ["Membre 3\n(Frontend / QA)",
         "UI + documentation",
         "Page /dashboard/services : statut Kong en temps réel\nAffichage liste transactions\nFix avertissement middleware.ts → proxy.ts",
         "Rapport Phase 5 + slides présentation\nTests manuels end-to-end\nCréer .env.example pour l'équipe"],
    ]
)

doc.add_paragraph()
doc.add_heading("Charge de travail estimée", level=2)
add_table(
    ["Tâche", "Complexité", "Durée", "Membre"],
    [
        ["docker-compose.yml Kong + DB Kong", "Haute", "1 jour", "Membre 1"],
        ["Microservice transactions CRUD", "Haute", "2 jours", "Membre 2"],
        ["Routes + plugin JWT Kong", "Haute", "1 jour", "Membre 1 + 2"],
        ["Page /dashboard/services", "Moyenne", "1 jour", "Membre 3"],
        ["Tests API (Jest/Supertest)", "Moyenne", "1 jour", "Membre 2 + 3"],
        ["Fix middleware → proxy.ts", "Faible", "2h", "Membre 3"],
        [".env.example pour l'équipe", "Faible", "30min", "Membre 3"],
        ["Déploiement Vercel + Railway", "Haute", "1 jour", "Membre 1"],
        ["Rapport Phase 5", "Moyenne", "0.5 jour", "Membre 3"],
    ]
)

doc.add_paragraph()
doc.add_heading("Règles de collaboration Git à respecter", level=2)
git_rules = [
    "Toujours créer une branche pour chaque tâche : git checkout -b feature/nom-tache",
    "Ne jamais pousser directement sur main sans PR",
    "Ouvrir une Pull Request sur GitHub et demander une revue à un autre membre",
    "Merger uniquement après approbation d'au moins 1 autre membre",
    "Synchroniser régulièrement : git pull origin main avant de commencer",
    "Écrire des messages de commit clairs : feat: / fix: / docs: / refactor:",
]
for rule in git_rules:
    p = doc.add_paragraph(style='List Number')
    p.add_run(rule)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  CONCLUSION
# ══════════════════════════════════════════════════════════════
doc.add_heading("Conclusion et appel à l'action", level=1)
doc.add_paragraph(
    "Ce rapport marque la fin de la Phase 4 et le lancement de la Phase 5. "
    "En 4 phases, le projet a construit une base solide et professionnelle : "
    "authentification double (Credentials + Google), interface utilisateur complète, "
    "base de données Prisma/PostgreSQL opérationnelle, et dépôt Git structuré avec 4 branches."
)
doc.add_paragraph(
    "La prochaine étape — Kong Gateway — est la plus importante techniquement. "
    "Elle différenciera ce projet d'une simple application CRUD et démontrera "
    "la maîtrise des architectures microservices."
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Action immédiate pour chaque membre :")
r.bold = True; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

actions = [
    "Récupérer le projet en local (Section 3 de ce rapport)",
    "Effectuer les 8 tests de validation (Section 5)",
    "Se coordonner avec les autres membres sur la répartition (Section 8)",
    "Créer sa branche Phase 5 et commencer le travail assigné",
]
for action in actions:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(action)
    r.bold = True

# ── Sauvegarde ──
output = "Rapport_Projet_TWM_Phase4.docx"
doc.save(output)
print(f"✅ Rapport généré : {output}")
