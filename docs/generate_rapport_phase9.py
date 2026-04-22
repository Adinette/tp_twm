from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Rapport_Projet_TWM_Phase9.docx'

PHASE9_BRANCH = 'feature/phase9-collaboration-release-hub'
PHASE9_BASE = 'feature/phase8-infra-cleanup'

RELEASE_TRACKS = [
    {
        'phase': 'Phase 7',
        'title': 'Correctif partage Phase 7',
        'branch': 'fix/phase7-reporting-resilience',
        'base': 'feature/phase7-microservices-advanced',
        'pr': 'PR #4',
        'url': 'https://github.com/Adinette/tp_twm/pull/4',
        'status': 'En revue',
        'usage': 'Equipe qui doit rester alignee sur la branche Phase 7 deja partagee.',
        'summary': 'Stabilise le reporting et la supervision sans reecrire la ligne de travail Phase 7.',
    },
    {
        'phase': 'Phase 8',
        'title': 'Consolidation complete Phase 8',
        'branch': 'feature/phase8-infra-cleanup',
        'base': 'feature/phase8-front-reporting-deployment-prep',
        'pr': 'PR #5',
        'url': 'https://github.com/Adinette/tp_twm/pull/5',
        'status': 'En revue',
        'usage': 'Collaborateurs qui veulent reprendre la derniere ligne front, reporting, order et infra.',
        'summary': 'Regroupe les correctifs order, reporting, services et les Dockerfiles restants.',
    },
]

PHASE9_DELIVERABLES = [
    ['app/lib/release-hub.ts', 'Nouveau', 'Source unique des branches, PR et commandes de recuperation.'],
    ['app/globals.css', 'Modifie', 'Nouvelle direction visuelle globale, surfaces, gradients et contraste.'],
    ['app/dashboard/layout.tsx', 'Modifie', 'Shell dashboard modernise avec ambiance visuelle plus lisible.'],
    ['app/components/DashboardSidebar.tsx', 'Modifie', 'Navigation groupee, acces rapide et resume livraison.'],
    ['app/dashboard/page.tsx', 'Modifie', 'Accueil dashboard repense autour des livraisons et des points d entree utiles.'],
    ['app/dashboard/versions/page.tsx', 'Nouveau', 'Centre de versions pour recuperer les bonnes branches et PR.'],
    ['app/page.tsx', 'Modifie', 'Landing page publique alignee sur la nouvelle logique de livraison.'],
    ['docs/generate_rapport_phase9.py', 'Nouveau', 'Generateur du rapport habituel pour la Phase 9.'],
]

VALIDATION_ROWS = [
    ['Lint TS/TSX', 'npx eslint app/dashboard/page.tsx app/dashboard/versions/page.tsx app/page.tsx app/components/DashboardSidebar.tsx app/lib/release-hub.ts app/dashboard/layout.tsx', 'OK'],
    ['Check cible', 'npx eslint app/dashboard/page.tsx --max-warnings=0', 'EXIT:0'],
    ['Diagnostics VS Code', 'get_errors sur les fichiers touches', 'OK sauf reliquat stale sur une suggestion Tailwind du dashboard'],
    ['DOCX', 'python docs/generate_rapport_phase9.py', 'OK - Rapport_Projet_TWM_Phase9.docx genere'],
]

PROBLEM_ROWS = [
    ['Branche Phase 7 deja partagee', 'Impossible de reecrire brutalement la ligne de travail d equipe', 'Creation d un hub de versions avec PR #4 et PR #5 pour orienter chaque collaborateur vers la bonne reprise'],
    ['UX trop diffuse', 'Le dashboard ne mettait pas assez en avant les livraisons et la bonne branche a recuperer', 'Refonte du shell, de l accueil dashboard et ajout d une page Versions dediee'],
    ['Windows + Turbopack instable', 'Certains services sont plus fiables avec Webpack dans cet environnement', 'Consigne explicite dans le hub et le rapport pour relancer les services critiques avec Webpack si necessaire'],
    ['Tailwind 4', 'Les classes arbitraires doivent suivre la syntaxe recommandee', 'Normalisation des classes et validation ESLint ciblee'],
]

PHASE9_HIGHLIGHTS = [
    'Refonte du dashboard pour orienter les collaborateurs vers les modules actifs et les branches a recuperer.',
    'Ajout d un centre de versions avec PR, branches, commandes Git et consignes de synchronisation.',
    'Amelioration visuelle du shell dashboard et de la landing page publique pour rendre les points d entree plus lisibles.',
]

COLLABORATION_CHECKLIST = [
    'Executer git fetch origin avant toute reprise locale.',
    'Choisir la bonne ligne: correctif Phase 7 partage ou consolidation complete Phase 8.',
    'Faire git switch puis git pull sur la branche cible avant toute nouvelle modification.',
    'Relancer ensuite les services utiles, en priorite avec Webpack quand Turbopack est instable sur Windows.',
]

doc = Document()


def setup_styles() -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_title_page() -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rapport de Projet - Phase 9')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run('Hub de versions collaboratif et refonte UI/UX du frontend principal')
    sub.font.size = Pt(15)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Projet : ').bold = True
    meta.add_run('tp_twm | SFMC Benin | Groupe 5\n')
    meta.add_run('Depot GitHub : ').bold = True
    meta.add_run('github.com/Adinette/tp_twm\n')
    meta.add_run('Branche active : ').bold = True
    meta.add_run(f'{PHASE9_BRANCH}\n')
    meta.add_run('Base technique : ').bold = True
    meta.add_run(f'{PHASE9_BASE}\n')
    meta.add_run('Date : ').bold = True
    meta.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()


def add_table(headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = value


def add_bullet(text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build_document() -> None:
    doc.add_heading('Table des matieres', level=1)
    toc_items = [
        '1. Objectif de la phase 9',
        '2. Branches et versions a recuperer',
        '3. Ce qui a ete livre en UI/UX',
        '4. Fichiers modifies',
        '5. Validation effectuee',
        '6. Problemes rencontres et solutions',
        '7. Mode operatoire collaborateur',
        '8. Suite recommandee',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading('1. Objectif de la phase 9', level=1)
    doc.add_paragraph(
        'La Phase 9 se concentre sur deux besoins immediats du projet : ameliorer la qualite percue du frontend principal '
        'et permettre a un collaborateur de recuperer sans ambiguite la bonne version partagee du travail. '
        'La solution retenue consiste a moderniser le shell du dashboard, clarifier la navigation et introduire un centre de versions '
        'qui expose les branches, PR et commandes Git utiles a la reprise.'
    )
    for item in PHASE9_HIGHLIGHTS:
        add_bullet(item)

    doc.add_paragraph()
    doc.add_heading('2. Branches et versions a recuperer', level=1)
    doc.add_paragraph(
        'Le point critique de la phase est la coexistence de plusieurs lignes utiles: la correction partagee de la Phase 7, '
        'la consolidation complete de la Phase 8 et la nouvelle branche Phase 9 qui documente et rend ces recuperations visibles dans le produit.'
    )
    add_table(
        ['Phase', 'Titre', 'Branche', 'Base', 'PR', 'Statut'],
        [
            [track['phase'], track['title'], track['branch'], track['base'], track['pr'], track['status']]
            for track in RELEASE_TRACKS
        ],
    )
    doc.add_paragraph('Resume collaborateur :')
    for track in RELEASE_TRACKS:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{track['phase']} - {track['title']} : ").bold = True
        paragraph.add_run(f"{track['summary']} Usage cible : {track['usage']} {track['pr']} ({track['url']}).")

    doc.add_paragraph('Branche de travail de la Phase 9 :')
    add_code(
        'git fetch origin\n'
        f'git switch {PHASE9_BRANCH}\n'
        f'git pull origin {PHASE9_BRANCH}'
    )

    doc.add_paragraph()
    doc.add_heading('3. Ce qui a ete livre en UI/UX', level=1)
    doc.add_heading('3.1 Dashboard principal', level=2)
    add_bullet('Refonte de la page dashboard pour mettre en avant la version recommandee, les PR actives et les acces rapides vers les modules utiles.')
    add_bullet('Modernisation du shell dashboard avec ambiance visuelle, surfaces plus lisibles et hierarchie plus claire.')
    add_bullet('Reorganisation de la sidebar avec navigation groupee et resume livraison.')

    doc.add_heading('3.2 Centre de versions', level=2)
    add_bullet('Ajout de /dashboard/versions comme point unique de reprise pour les collaborateurs.')
    add_bullet('Chaque carte expose branche, base, PR, public cible, resume et commandes Git a copier.')
    add_bullet('Ajout d une checklist de synchronisation pour limiter les erreurs de reprise locale.')

    doc.add_heading('3.3 Accueil public', level=2)
    add_bullet('Refonte de la landing page publique pour presenter le projet comme une plateforme livrable et non seulement comme une demo technique.')
    add_bullet('Mise en avant des lignes partagees Phase 7 et Phase 8 pour faciliter la communication equipe.')

    doc.add_paragraph()
    doc.add_heading('4. Fichiers modifies', level=1)
    add_table(['Fichier', 'Statut', 'Apport'], PHASE9_DELIVERABLES)

    doc.add_paragraph()
    doc.add_heading('5. Validation effectuee', level=1)
    doc.add_paragraph('Validation ciblee realisee sur les nouveaux ecrans et leurs metadonnees :')
    add_table(['Type', 'Commande / outil', 'Resultat'], VALIDATION_ROWS)
    doc.add_paragraph('Commandes utiles cote collaborateur pour valider rapidement la livraison :')
    add_code(
        'npx eslint app/dashboard/page.tsx app/dashboard/versions/page.tsx app/page.tsx '
        'app/components/DashboardSidebar.tsx app/lib/release-hub.ts app/dashboard/layout.tsx\n'
        'npm run dev'
    )

    doc.add_paragraph()
    doc.add_heading('6. Problemes rencontres et solutions', level=1)
    add_table(['Probleme', 'Cause', 'Solution'], PROBLEM_ROWS)

    doc.add_paragraph()
    doc.add_heading('7. Mode operatoire collaborateur', level=1)
    doc.add_paragraph('La page /dashboard/versions et le rapport donnent le meme protocole de reprise :')
    for item in COLLABORATION_CHECKLIST:
        add_bullet(item)
    doc.add_paragraph('Exemple de sequence pour reprendre une version partagee :')
    add_code(
        'git fetch origin\n'
        'git switch fix/phase7-reporting-resilience\n'
        'git pull origin fix/phase7-reporting-resilience\n\n'
        'git switch feature/phase8-infra-cleanup\n'
        'git pull origin feature/phase8-infra-cleanup'
    )

    doc.add_paragraph()
    doc.add_heading('8. Suite recommandee', level=1)
    add_bullet('Ouvrir la future PR de la branche Phase 9 une fois la revue visuelle terminee.')
    add_bullet('Ajouter quelques captures d ecran des nouveaux ecrans si le livrable final le demande.')
    add_bullet('Valider en equipe le choix entre la ligne Phase 7 partagee et la ligne Phase 8 consolidee avant toute nouvelle phase technique.')
    add_bullet('Conserver le centre de versions comme point d entree pour les reprises futures ou pour la demonstration finale.')

    doc.add_paragraph()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'La Phase 9 ne cherche pas a ajouter un nouveau microservice. Elle rend le projet plus exploitable par l equipe : '
        'le frontend est plus clair, la reprise des branches est explicite et la livraison des correctifs precedents devient comprehensible. '
        'Le resultat est une couche de coordination concrete entre la technique, le design et le travail collaboratif.'
    )


def main() -> None:
    setup_styles()
    add_title_page()
    build_document()
    doc.save(OUTPUT)
    print(f'Rapport genere : {OUTPUT}')


if __name__ == '__main__':
    main()